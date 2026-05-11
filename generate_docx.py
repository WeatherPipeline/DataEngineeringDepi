import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_DIR = r"C:\Users\DELL\OneDrive\Desktop\P\docs"
OUTPUT = r"C:\Users\DELL\OneDrive\Desktop\P\Weather_ETL_Pipeline_Documentation.docx"


def read_md(filename):
    path = os.path.join(DOC_DIR, filename)
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def parse_md_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith("|") and not set(line.replace("|", "").replace("-", "").replace(" ", "")) == set():
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if cells:
                rows.append(cells)
    if len(rows) >= 2:
        return rows[0], rows[2:]
    return None, None


def is_table_line(line):
    return line.strip().startswith("|") and "---" not in line.strip().replace(" ", "")[:10]


def is_separator_line(line):
    stripped = line.strip()
    return stripped.startswith("|") and all(c in "|- " for c in stripped)


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    pPr.append(ind)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h.strip()
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(cell, "D9E2F3")
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val.strip()
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def add_md_content(doc, md_text, max_heading_level=2):
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, "\n".join(code_lines))
            i += 1
            continue

        if is_table_line(line):
            table_lines = []
            while i < len(lines) and (is_table_line(lines[i]) or is_separator_line(lines[i])):
                table_lines.append(lines[i])
                i += 1
            headers, rows = parse_md_table(table_lines)
            if headers and rows:
                add_table(doc, headers, rows)
            continue

        if line.startswith("# ") and max_heading_level >= 1:
            text = line.lstrip("#").strip()
            doc.add_heading(text, level=1)
            i += 1
            continue

        if line.startswith("## ") and max_heading_level >= 2:
            text = line.lstrip("#").strip()
            doc.add_heading(text, level=2)
            i += 1
            continue

        if line.startswith("### ") and max_heading_level >= 3:
            text = line.lstrip("#").strip()
            doc.add_heading(text, level=3)
            i += 1
            continue

        if line.startswith("- ") or line.startswith("* "):
            text = line.lstrip("-*").strip()
            p = doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        if line.strip().startswith(tuple(f"{n}." for n in range(1, 20))):
            text = line.strip()
            num_end = text.index(".")
            content = text[num_end + 1:].strip()
            p = doc.add_paragraph(content, style="List Number")
            i += 1
            continue

        p = doc.add_paragraph(line.strip())
        i += 1


def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # ========== COVER PAGE ==========
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Weather ETL Pipeline\nfor Egypt Governorates")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    doc.add_paragraph()
    add_horizontal_line(doc)
    doc.add_paragraph()

    info_items = [
        ("Project", "Weather ETL Pipeline"),
        ("Version", "1.0"),
        ("Author", "Mahmoud"),
        ("Date", "May 2026"),
        ("Status", "Operational"),
        ("Classification", "Internal"),
        ("Technology", "Apache Airflow 3.2.0 / Python / SQL Server"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.size = Pt(12)
        run_value = p.add_run(value)
        run_value.font.size = Pt(12)

    doc.add_page_break()

    # ========== DOCUMENT CONTROL ==========
    doc.add_heading("Document Control", level=1)

    doc.add_heading("Version History", level=2)
    add_table(doc,
        ["Version", "Date", "Author", "Description"],
        [
            ["0.1", "April 2026", "Mahmoud", "Initial project setup with Airflow 2.8.1"],
            ["0.5", "April 2026", "Mahmoud", "Migrated to Airflow 3.2.0, added scikit-learn"],
            ["0.9", "May 7, 2026", "Mahmoud", "DAG finalized, hourly schedule active"],
            ["1.0", "May 10, 2026", "Mahmoud", "Production-ready, full documentation"],
        ]
    )

    doc.add_heading("Reviewers", level=2)
    add_table(doc,
        ["Role", "Name", "Date"],
        [
            ["Developer", "Mahmoud", "May 2026"],
        ]
    )

    doc.add_page_break()

    # ========== TABLE OF CONTENTS ==========
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Introduction",
        "   2.1 Purpose",
        "   2.2 Intended Audience",
        "   2.3 Project Background",
        "   2.4 Scope",
        "3. Project Overview",
        "   3.1 Problem Statement",
        "   3.2 Solution Description",
        "   3.3 Goals and Objectives",
        "   3.4 Stakeholders",
        "4. System Requirements",
        "   4.1 Functional Requirements",
        "   4.2 Non-Functional Requirements",
        "   4.3 Technical Requirements",
        "   4.4 Infrastructure Requirements",
        "5. Architecture and Design",
        "   5.1 High-Level Architecture",
        "   5.2 Component Design",
        "   5.3 Network Topology",
        "   5.4 Data Flow",
        "   5.5 Volume Mounts",
        "6. Technology Stack",
        "7. ETL Pipeline",
        "   7.1 DAG Configuration",
        "   7.2 Task 1: create_weather_table",
        "   7.3 Task 2: load_weather_data",
        "   7.4 Deduplication Mechanism",
        "8. Database Documentation",
        "   8.1 Schema",
        "   8.2 Constraints and Indexes",
        "   8.3 Views",
        "   8.4 Common Queries",
        "9. Data Source: Open-Meteo API",
        "10. Deployment Documentation",
        "   10.1 Deployment Architecture",
        "   10.2 Environment Details",
        "   10.3 Installation Steps",
        "   10.4 Configuration Changes",
        "   10.5 Backup and Recovery",
        "   10.6 Troubleshooting",
        "11. Development Standards",
        "12. Operational Runbook",
        "13. Glossary",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        if not item.startswith("   "):
            for run in p.runs:
                run.bold = True

    doc.add_page_break()

    # ========== EXECUTIVE SUMMARY ==========
    doc.add_heading("1. Executive Summary", level=1)

    doc.add_heading("Project Overview", level=2)
    doc.add_paragraph(
        "The Weather ETL Pipeline is a data engineering project that automates the collection, "
        "transformation, and storage of real-time weather data for all 27 Egyptian governorates. "
        "The system runs on an hourly schedule using Apache Airflow 3.2.0, fetches data from the "
        "free Open-Meteo API, applies temperature alert classification and linear regression "
        "prediction, and stores results in Microsoft SQL Server."
    )

    doc.add_heading("Business Objective", level=2)
    doc.add_paragraph(
        "To establish a reliable, automated weather data collection system that enables historical "
        "trend analysis, temperature alert monitoring, and predictive temperature modeling for "
        "Egyptian governorates, supporting data-driven decision making."
    )

    doc.add_heading("Expected Outcome", level=2)
    outcomes = [
        "Hourly weather observations for 27 Egyptian cities stored in a structured database",
        "Automated alert classification for extreme temperatures (High >36C, Low <15C)",
        "ML-predicted temperature values based on humidity and wind speed correlations",
        "Deduplicated historical dataset growing at approximately 648 records per day",
    ]
    for o in outcomes:
        doc.add_paragraph(o, style="List Bullet")

    doc.add_heading("High-Level Scope", level=2)
    doc.add_paragraph(
        "Single-machine deployment on a Windows workstation. Docker-based Airflow orchestration "
        "with CeleryExecutor. SQL Server target database. No cloud infrastructure, no external "
        "authentication, no user-facing application."
    )

    doc.add_page_break()

    # ========== INTRODUCTION ==========
    doc.add_heading("2. Introduction", level=1)

    doc.add_heading("2.1 Purpose", level=2)
    doc.add_paragraph(
        "This document serves as the comprehensive technical reference for the Weather ETL Pipeline "
        "project. It covers system architecture, data pipeline design, database schema, deployment "
        "procedures, technology decisions, development standards, and operational runbooks. It is "
        "intended to be the single source of truth for understanding, deploying, and maintaining "
        "the system."
    )

    doc.add_heading("2.2 Intended Audience", level=2)
    add_table(doc,
        ["Audience", "Use Case"],
        [
            ["Data Engineers", "Understand pipeline logic, extend functionality"],
            ["DevOps / Infrastructure", "Deploy, configure, troubleshoot"],
            ["Project Stakeholders", "Review system capabilities and requirements"],
            ["New Team Members", "Onboarding and knowledge transfer"],
        ]
    )

    doc.add_heading("2.3 Project Background", level=2)
    doc.add_paragraph(
        "The project was initiated in April 2026 as a data engineering exercise to build an "
        "automated ETL pipeline for Egyptian weather data. It evolved through several iterations: "
        "starting with Airflow 2.8.1 and a basic DAG, then migrating to Airflow 3.2.0, adding "
        "scikit-learn for ML predictions, and stabilizing on the current architecture with "
        "CeleryExecutor, custom Docker images, and SQL Server integration."
    )

    doc.add_heading("2.4 Scope", level=2)
    doc.add_heading("In-Scope", level=3)
    in_scope = [
        "Weather data extraction from Open-Meteo API for 27 Egyptian governorates",
        "Hourly automated scheduling via Apache Airflow",
        "Temperature alert classification (High / Normal / Low)",
        "Linear regression temperature prediction model",
        "Data storage in SQL Server with deduplication",
        "Docker-based deployment with CeleryExecutor",
        "Operational runbook and troubleshooting guides",
    ]
    for item in in_scope:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Out-of-Scope", level=3)
    out_scope = [
        "Cloud deployment (AWS, Azure, GCP)",
        "Real-time streaming (Kafka, Spark Streaming)",
        "User-facing dashboard or reporting UI",
        "Email or Slack alerting notifications",
        "Multi-node or clustered deployment",
        "Data archival or retention policies",
        "CI/CD pipeline for automated deployment",
    ]
    for item in out_scope:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ========== PROJECT OVERVIEW ==========
    doc.add_heading("3. Project Overview", level=1)

    doc.add_heading("3.1 Problem Statement", level=2)
    doc.add_paragraph(
        "There was no automated system to collect and store weather observations for Egyptian "
        "governorates in a structured, queryable format. Weather data was not available for "
        "historical analysis, trend detection, or alert monitoring."
    )

    doc.add_heading("3.2 Solution Description", level=2)
    doc.add_paragraph(
        "An hourly batch ETL pipeline built on Apache Airflow that fetches current weather "
        "conditions (temperature, humidity, wind speed) from the free Open-Meteo API, classifies "
        "temperature alerts, predicts temperature using a linear regression model trained on each "
        "batch, and stores the results in SQL Server with deduplication by city and timestamp."
    )

    doc.add_heading("3.3 Goals and Objectives", level=2)
    add_table(doc,
        ["Goal", "Objective", "Status"],
        [
            ["Automated data collection", "Hourly weather fetch for 27 cities", "Achieved"],
            ["Data quality", "Deduplication and constraint enforcement", "Achieved"],
            ["Alert system", "Temperature classification per observation", "Achieved"],
            ["ML prediction", "LinearRegression per batch", "Achieved"],
            ["Operational reliability", "Auto-restart, retry logic, health checks", "Achieved"],
        ]
    )

    doc.add_heading("3.4 Stakeholders", level=2)
    add_table(doc,
        ["Role", "Name", "Responsibility"],
        [
            ["Developer / Owner", "Mahmoud", "Design, implementation, documentation, operations"],
        ]
    )

    doc.add_page_break()

    # ========== SYSTEM REQUIREMENTS ==========
    doc.add_heading("4. System Requirements", level=1)
    doc.add_paragraph("Full requirements are documented in docs/requirements.md. Summary below.")
    doc.add_paragraph()

    add_md_content(doc, read_md("requirements.md"), max_heading_level=2)

    doc.add_page_break()

    # ========== ARCHITECTURE ==========
    doc.add_heading("5. Architecture and Design", level=1)
    doc.add_paragraph("Full architecture documentation is in docs/architecture.md.")
    doc.add_paragraph()

    add_md_content(doc, read_md("architecture.md"), max_heading_level=2)

    doc.add_page_break()

    # ========== TECHNOLOGY STACK ==========
    doc.add_heading("6. Technology Stack", level=1)
    doc.add_paragraph("Full technology breakdown is in docs/tech-stack.md.")
    doc.add_paragraph()

    add_md_content(doc, read_md("tech-stack.md"), max_heading_level=2)

    doc.add_page_break()

    # ========== ETL PIPELINE ==========
    doc.add_heading("7. ETL Pipeline", level=1)
    doc.add_paragraph("Full pipeline documentation is in docs/pipeline.md.")
    doc.add_paragraph()

    add_md_content(doc, read_md("pipeline.md"), max_heading_level=2)

    doc.add_page_break()

    # ========== DATABASE ==========
    doc.add_heading("8. Database Documentation", level=1)
    doc.add_paragraph("Full database documentation is in docs/database.md.")
    doc.add_paragraph()

    add_md_content(doc, read_md("database.md"), max_heading_level=2)

    doc.add_page_break()

    # ========== DATA SOURCE ==========
    doc.add_heading("9. Data Source: Open-Meteo API", level=1)
    doc.add_paragraph("Full API documentation is in docs/data-source.md.")
    doc.add_paragraph()

    add_md_content(doc, read_md("data-source.md"), max_heading_level=2)

    doc.add_page_break()

    # ========== DEPLOYMENT ==========
    doc.add_heading("10. Deployment Documentation", level=1)
    doc.add_paragraph("Full deployment documentation is in docs/deployment.md.")
    doc.add_paragraph()

    add_md_content(doc, read_md("deployment.md"), max_heading_level=2)

    doc.add_page_break()

    # ========== DEVELOPMENT STANDARDS ==========
    doc.add_heading("11. Development Standards", level=1)
    doc.add_paragraph("Full development standards are in docs/development-standards.md.")
    doc.add_paragraph()

    add_md_content(doc, read_md("development-standards.md"), max_heading_level=2)

    doc.add_page_break()

    # ========== RUNBOOK ==========
    doc.add_heading("12. Operational Runbook", level=1)
    doc.add_paragraph("Full runbook is in docs/runbook.md.")
    doc.add_paragraph()

    add_md_content(doc, read_md("runbook.md"), max_heading_level=2)

    doc.add_page_break()

    # ========== GLOSSARY ==========
    doc.add_heading("13. Glossary", level=1)
    add_table(doc,
        ["Term", "Definition"],
        [
            ["ETL", "Extract, Transform, Load -- the process of moving data from source to target with transformations"],
            ["DAG", "Directed Acyclic Graph -- Airflow's representation of a workflow with task dependencies"],
            ["CeleryExecutor", "Airflow executor that distributes tasks to Celery workers via a message broker"],
            ["WSL2", "Windows Subsystem for Linux version 2 -- used by Docker Desktop as the virtualization backend"],
            ["ODBC", "Open Database Connectivity -- a standard API for accessing database management systems"],
            ["pyodbc", "Python module that provides ODBC database connectivity"],
            ["SQLAlchemy", "Python SQL toolkit and Object-Relational Mapping library"],
            ["Fernet Key", "Symmetric encryption key used by Airflow to encrypt connection passwords and variables"],
            ["Celery", "Distributed task queue library used by Airflow's CeleryExecutor"],
            ["Redis", "In-memory data store used as the Celery message broker"],
            ["Celery Worker", "Process that executes tasks from the Celery queue"],
            ["Airflow Scheduler", "Process that evaluates DAG schedules and queues task instances for execution"],
            ["Dag Processor", "Process that parses DAG files and registers them in the metadata database"],
            ["Triggerer", "Airflow process that manages asynchronous triggers for deferrable operators"],
            ["Bind Mount", "Docker volume type that maps a host directory into a container"],
            ["Named Volume", "Docker volume managed by Docker, stored in Docker's internal storage"],
            ["host.docker.internal", "Special DNS name in Docker Desktop that resolves to the host machine's IP"],
            ["Open-Meteo", "Free, open-source weather API that provides forecast and historical weather data"],
            ["Linear Regression", "Statistical method that models the relationship between a dependent variable and one or more independent variables"],
            ["scikit-learn", "Python machine learning library providing classification, regression, and clustering algorithms"],
            ["pandas", "Python data manipulation library providing DataFrame data structures"],
            ["Idempotent", "An operation that produces the same result whether executed once or multiple times"],
            ["Catchup", "Airflow behavior that backfills missed DAG runs between the start date and the current date"],
            ["Fernet", "Symmetric encryption implementation in the cryptography library, used by Airflow for secrets"],
            ["Health Check", "Docker mechanism to verify that a container is functioning correctly"],
        ]
    )

    # ========== SIGN-OFF ==========
    doc.add_page_break()
    doc.add_heading("Sign-Off", level=1)
    doc.add_paragraph()
    add_table(doc,
        ["Role", "Name", "Signature", "Date"],
        [
            ["Developer", "Mahmoud", "", "May 2026"],
        ]
    )

    doc.save(OUTPUT)
    print(f"Document saved to: {OUTPUT}")


if __name__ == "__main__":
    build_document()
